#!/usr/bin/env python3
"""Generate a single complex-structure .docx for manual / parser testing only.

Contains: header, footer, TOC field (with placeholder result), headings, ~2 pages of body
text (with manual page break), body image, **figure caption + table caption** (Word Caption /
题注 style — distinct from TOC), multiple **visibly styled** external hyperlinks
(blue #0563C1 + underline), and a table with a hyperlink in a cell.

Not imported by application code. Run from repo:

  cd document-service
  .venv\\Scripts\\python.exe scripts/generate_complex_docx_fixture.py -o tests/fixtures/complex-structure-fixture.docx

Open the file in Word and press F9 (Update Fields) to refresh the table of contents
if you need a fully materialized TOC; the embedded placeholder text is still useful
for offline parsing tests.
"""

from __future__ import annotations

import argparse
import base64
import io
import sys
from pathlib import Path

# Minimal valid 1x1 PNG (grey pixel)
_MINI_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)

# Filler text to stretch toward ~2 pages in Word (print layout); tune if needed.
_PAGE_FILLER_PARA = (
    "本段为占位正文，用于将 Fixture 文档撑开至约两页，便于检查页眉页脚、分页与目录。"
    "解析管线应能抽取连续段落；超链接在下方专节中重复出现以便肉眼辨认。"
    "若版面仍不足两页，可在 Word 中调大字号或行距验证分页。"
)


def _add_external_hyperlink(paragraph: object, url: str, text: str) -> None:
    """Append w:hyperlink with visible link styling (blue + underline) like Word default."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    part = paragraph.part  # type: ignore[attr-defined]
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    # Explicit color + underline so links are visible even if template lacks "Hyperlink" style.
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    h.append(r)
    paragraph._p.append(h)  # type: ignore[attr-defined]


def _apply_caption_style(document: object, paragraph: object) -> str:
    """Use built-in caption style (en: Caption, zh: 题注) so it is not tagged as TOC in parse."""
    for name in ("Caption", "题注"):
        try:
            paragraph.style = document.styles[name]  # type: ignore[index]
            return name
        except KeyError:
            continue
    return ""


def _append_manual_page_break(document: object) -> None:
    from docx.enum.text import WD_BREAK

    p = document.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _append_toc_field_paragraph(document: object) -> None:
    """Insert a w:p containing a TOC field (levels 1–3) with a visible placeholder result."""
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml.parser import parse_xml

    # Word TOC instruction; separate/run result lets parsers see text without opening Word.
    xml = (
        "<w:p {ns}>"
        "<w:r><w:fldChar w:fldCharType=\"begin\"/></w:r>"
        "<w:r><w:instrText xml:space=\"preserve\"> TOC \\o &quot;1-3&quot; \\h \\z \\u </w:instrText></w:r>"
        "<w:r><w:fldChar w:fldCharType=\"separate\"/></w:r>"
        "<w:r><w:t>（解析测试占位：请在 Word 中 F9 更新域以刷新目录）</w:t></w:r>"
        "<w:r><w:fldChar w:fldCharType=\"end\"/></w:r>"
        "</w:p>"
    ).format(ns=nsdecls("w"))
    p = parse_xml(xml)
    body = document.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(p)
    else:
        body.append(p)


def generate(output_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt

    doc = Document()

    # --- Header / footer (first section) ---
    sec = doc.sections[0]
    hdr = sec.header
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.text = "页眉-HDR-FIXTURE-001 | 复杂结构测试"
    hp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    fp.text = "页脚-FTR-FIXTURE-002"
    fp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Body: first paragraph as title (some environments have no default empty para) ---
    if len(doc.paragraphs) < 1:
        doc.add_paragraph("")
    t = doc.paragraphs[0]
    t.text = "复杂结构 Fixture 文档"
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    try:
        t.runs[0].font.size = Pt(18)
        t.runs[0].bold = True
    except IndexError:
        pass

    toc_heading = doc.add_paragraph()
    toc_heading.add_run("目录").bold = True

    # Body order: [0] title, [1] "目录" label → insert TOC field at index 2.
    _append_toc_field_paragraph(doc)

    doc.add_paragraph("")  # spacer

    # ----- 第一页：超链接（蓝字+下划线）、图片、长正文撑版心 -----
    doc.add_heading("一、图片与超链接（行内链接）", level=1)
    doc.add_paragraph(
        "说明：以下「蓝色 + 单下划线」文本为外部超链接；若仍为黑色，请确认用 Word 桌面版打开且未禁用「显示域颜色」。"
    )
    p = doc.add_paragraph()
    p.add_run("正文前缀（黑） ")
    _add_external_hyperlink(p, "https://example.com/fixture", "EX-LINK-超链-蓝字")
    p.add_run(" 正文后缀（黑）。")
    p.add_run().add_picture(io.BytesIO(_MINI_PNG), width=Inches(0.9))
    p.add_run(" 图后字（黑）")

    # Block figure + caption (Caption / 题注) — parser should keep these; not `sourceRegion: toc`.
    fig_block = doc.add_paragraph()
    fig_block.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fig_block.add_run().add_picture(io.BytesIO(_MINI_PNG), width=Inches(1.1))
    cap_fig = doc.add_paragraph(
        "图 1 Fixture 专用示例图注（使用 Caption/题注 样式，与目录样式 TOC 1 等区分）"
    )
    _apply_caption_style(doc, cap_fig)

    p_only = doc.add_paragraph()
    _add_external_hyperlink(
        p_only,
        "https://www.example.org/path?fixture=1",
        "【整段仅超链接】整行应为蓝色下划线，便于肉眼辨认",
    )

    doc.add_heading("超链接专节（多条独立链接）", level=2)
    for label, url, anchor in (
        ("Python", "https://www.python.org/", "python.org 首页"),
        ("Wikipedia", "https://www.wikipedia.org/", "维基百科"),
        ("RFC2606", "https://www.rfc-editor.org/rfc/rfc2606.html", "保留域名说明 RFC"),
    ):
        pl = doc.add_paragraph()
        pl.add_run(f"{label}：")
        _add_external_hyperlink(pl, url, anchor)

    doc.add_paragraph("以下为重复占位段落，用于将第一页撑满（约一页后再手动分页）。")
    for i in range(12):
        doc.add_paragraph(f"（第一页填充 {i + 1:02d}/12）{_PAGE_FILLER_PARA}")

    _append_manual_page_break(doc)

    # ----- 第二页：表格内超链、结语 -----
    doc.add_heading("二、表格与单元格超链（第二页）", level=1)
    doc.add_paragraph("下表「左下」单元格中含超链接，应为蓝色下划线。")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "表头-A"
    table.cell(0, 1).text = "表头-B"
    c = table.cell(1, 0)
    if len(c.paragraphs) < 1:
        c.add_paragraph("")
    cp = c.paragraphs[0]
    cp.text = ""
    cp.add_run("单元格前缀 ")
    _add_external_hyperlink(cp, "https://example.com/cell", "CELL-HLINK-蓝字")
    table.cell(1, 1).text = "单元格-纯文本"

    cap_tbl = doc.add_paragraph("表 1 Fixture 表格题注示例（同为 Caption/题注 样式）")
    _apply_caption_style(doc, cap_tbl)

    doc.add_heading("三、结语与补充段落", level=2)
    doc.add_paragraph("结语：目录层级与收尾。以下为第二页补充占位。")
    for j in range(8):
        doc.add_paragraph(f"（第二页填充 {j + 1:02d}/8）{_PAGE_FILLER_PARA}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> int:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument(
        "-o",
        "--output",
        type=Path,
        default=Path(__file__).resolve().parent.parent / "tests" / "fixtures" / "complex-structure-fixture.docx",
        help="Output .docx path (default: tests/fixtures/complex-structure-fixture.docx)",
    )
    args = p.parse_args()
    try:
        generate(args.output)
    except Exception as e:
        print(f"Failed: {e}", file=sys.stderr)
        return 1
    print(f"Wrote: {args.output.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
