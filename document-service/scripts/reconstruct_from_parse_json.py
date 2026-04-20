#!/usr/bin/env python3
"""Reconstruct artifacts from document-service parse JSON (lossy).

Uses **parseDocument.pages** block tree (titles, paragraphs, figures, tables, lists) when
building structured outputs; **embeddedImagePayloads** supplies figure bytes (``assetRef`` /
``imageKey``).

Outputs
-------
- **Structured** (``--formats structured`` or ``all``): one primary file by ``fileExt`` —
  ``rebuilt.docx`` (doc/docx), ``rebuilt.pdf`` (pdf), ``rebuilt.xlsx`` (xls/xlsx). Not a byte-
  for-byte restore of the original upload; semantic re-export only.
- **plain / markdown / html**: ``chunkView``-based (tables as summary rows); **assets/** from
  embedded payloads.

This is **not** ML vector ``embedding`` — only base64 image payloads.

Examples::

    cd document-service
    .\\.venv\\Scripts\\python.exe scripts\\reconstruct_from_parse_json.py ^
      -i D:\\tmp\\parse-response.json -o D:\\tmp\\reconstructed --formats all

    .\\.venv\\Scripts\\python.exe scripts\\reconstruct_from_parse_json.py ^
      -i doc-only.json --embedded-json payloads.json -o out\\ --formats structured
"""

from __future__ import annotations

import argparse
import base64
import binascii
import html
import json
import mimetypes
import sys
from io import BytesIO
from pathlib import Path
from typing import Any

# Allow ``from app...`` when run as ``python scripts/reconstruct_from_parse_json.py``
_SERVICE_ROOT = Path(__file__).resolve().parents[1]
if str(_SERVICE_ROOT) not in sys.path:
    sys.path.insert(0, str(_SERVICE_ROOT))

from app.parsers.unified_builders import build_chunk_view, raw_text_from_pages  # noqa: E402


def _plain_text_from_spans(spans: list) -> str:
    return "".join(str(s.get("text") or "") for s in spans if isinstance(s, dict))


def _plain_text_from_lines(lines: list[dict] | None) -> str:
    if not lines:
        return ""
    parts: list[str] = []
    for line in lines:
        if not isinstance(line, dict):
            continue
        parts.append(_plain_text_from_spans(line.get("spans") or []))
    return "\n".join(p for p in parts if p)


def _sorted_pages(pages: list[dict]) -> list[dict]:
    return sorted(
        [p for p in pages if isinstance(p, dict)],
        key=lambda p: int(p.get("pageIndex") or 0),
    )


def _sorted_blocks(blocks: list) -> list[dict]:
    return sorted(
        [b for b in blocks if isinstance(b, dict)],
        key=lambda b: int(b.get("readingOrder") if b.get("readingOrder") is not None else 0),
    )


def _load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def _parse_document(root: dict[str, Any]) -> dict[str, Any]:
    pd = root.get("parseDocument")
    if isinstance(pd, dict):
        return pd
    return root


def _embedded_map(root: dict[str, Any]) -> dict[str, dict[str, Any]]:
    emb = root.get("embeddedImagePayloads")
    if isinstance(emb, dict):
        return {str(k): v for k, v in emb.items() if isinstance(v, dict)}
    return {}


def _guess_ext(content_type: str) -> str:
    ct = (content_type or "").split(";")[0].strip().lower()
    ext = mimetypes.guess_extension(ct or "")
    if ext == ".jpe":
        ext = ".jpg"
    if ext:
        return ext
    if "png" in ct:
        return ".png"
    if "jpeg" in ct or "jpg" in ct:
        return ".jpg"
    if "gif" in ct:
        return ".gif"
    if "webp" in ct:
        return ".webp"
    return ".bin"


def _ext_from_image_magic(data: bytes) -> str | None:
    """Pick file extension from binary signature (more reliable than ``contentType`` alone)."""
    if len(data) < 12:
        return None
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return ".png"
    if data[:3] == b"\xff\xd8\xff":
        return ".jpg"
    if data[:6] in (b"GIF87a", b"GIF89a"):
        return ".gif"
    if data[:4] == b"RIFF" and len(data) >= 12 and data[8:12] == b"WEBP":
        return ".webp"
    if data[:2] == b"BM":
        return ".bmp"
    return None


def _decode_payload_row(row: dict[str, Any]) -> tuple[bytes, str] | None:
    b64 = str(row.get("base64") or "").strip()
    if not b64:
        return None
    try:
        raw = base64.b64decode(b64, validate=False)
    except (binascii.Error, ValueError):
        return None
    if not raw:
        return None
    ct = str(row.get("contentType") or "application/octet-stream")
    return raw, ct


def extract_images(
    embedded: dict[str, dict[str, Any]],
    assets_dir: Path,
) -> dict[str, str]:
    """Write image files; return map imageKey -> relative path under output root."""
    assets_dir.mkdir(parents=True, exist_ok=True)
    rel: dict[str, str] = {}
    for key, row in embedded.items():
        if not key:
            continue
        dec = _decode_payload_row(row)
        if not dec:
            continue
        raw, ct = dec
        # Prefer magic bytes: wrong ``contentType`` + ``.png`` name breaks many viewers while Word still opens the blob.
        ext = _ext_from_image_magic(raw) or _guess_ext(ct)
        safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in key)
        fname = f"{safe}{ext}"
        path = assets_dir / fname
        if path.exists():
            stem, suf = path.stem, path.suffix
            n = 2
            while path.exists():
                path = assets_dir / f"{stem}_{n}{suf}"
                n += 1
        path.write_bytes(raw)
        rel[key] = str(Path(assets_dir.name) / path.name)
    return rel


def build_image_bytes_from_embedded(
    embedded: dict[str, dict[str, Any]],
) -> dict[str, tuple[bytes, str]]:
    """Map imageKey -> (raw bytes, content-type) for in-document embedding."""
    out: dict[str, tuple[bytes, str]] = {}
    for key, row in embedded.items():
        dec = _decode_payload_row(row)
        if not dec:
            continue
        out[str(key)] = dec
    return out


def _find_inner_table_with_rows(block: dict[str, Any]) -> dict[str, Any] | None:
    for ib in block.get("blocks") or []:
        if isinstance(ib, dict) and str(ib.get("type") or "") == "table" and ib.get("rows"):
            return ib
    return None


def _table_rows_to_matrix(table_body: dict[str, Any]) -> list[list[str]]:
    rows_out: list[list[str]] = []
    for row in table_body.get("rows") or []:
        if not isinstance(row, dict):
            continue
        cells = row.get("cells") or []
        line: list[str] = []
        for c in cells:
            if isinstance(c, dict):
                line.append(_plain_text_from_spans(c.get("spans") or []).strip())
            else:
                line.append("")
        rows_out.append(line)
    return rows_out


def _docx_add_list_paragraph(doc: Any, text: str, *, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    try:
        doc.add_paragraph(text, style=style)
    except (KeyError, ValueError):
        prefix = f"{len(doc.paragraphs)}." if numbered else "•"
        doc.add_paragraph(f"{prefix} {text}")


def emit_docx_from_pages(
    pages: list[dict[str, Any]],
    image_bytes: dict[str, tuple[bytes, str]],
    out_path: Path,
) -> None:
    from docx import Document  # type: ignore
    from docx.shared import Inches  # type: ignore

    doc = Document()

    def walk(blocks: list, *, list_numbered: bool | None) -> None:
        for block in _sorted_blocks(blocks):
            btype = str(block.get("type") or "")
            if btype == "title":
                t = _plain_text_from_lines(block.get("lines")).strip()
                if not t:
                    continue
                level = max(1, min(9, int(block.get("level") or 1)))
                doc.add_heading(t, level=min(level, 9))
            elif btype == "paragraph":
                t = _plain_text_from_lines(block.get("lines")).strip()
                if t:
                    doc.add_paragraph(t)
            elif btype in ("formula", "code"):
                t = _plain_text_from_lines(block.get("lines")).strip()
                if not t:
                    t = str(block.get("formulaLatex") or "").strip()
                if t:
                    doc.add_paragraph(t)
            elif btype == "list":
                style = str(block.get("listStyle") or "bullet")
                num = style.lower() in ("number", "ordered", "decimal")
                walk(block.get("blocks") or [], list_numbered=num)
            elif btype == "list_item":
                t = _plain_text_from_lines(block.get("lines")).strip()
                if t:
                    _docx_add_list_paragraph(doc, t, numbered=bool(list_numbered))
                walk(block.get("blocks") or [], list_numbered=list_numbered)
            elif btype == "figure":
                ref = str(block.get("assetRef") or "").strip()
                if ref and ref in image_bytes:
                    raw, _ct = image_bytes[ref]
                    try:
                        doc.add_picture(BytesIO(raw), width=Inches(5.5))
                    except Exception:
                        doc.add_paragraph(f"[figure decode failed: {ref}]")
                elif ref:
                    doc.add_paragraph(f"[missing image payload: {ref}]")
                walk(block.get("blocks") or [], list_numbered=None)
            elif btype == "table":
                inner_blocks = block.get("blocks") or []
                for ib in _sorted_blocks(inner_blocks):
                    if not isinstance(ib, dict):
                        continue
                    it = str(ib.get("type") or "")
                    if it == "paragraph":
                        cap = _plain_text_from_lines(ib.get("lines")).strip()
                        if cap:
                            doc.add_paragraph(cap)
                    elif it == "table" and ib.get("rows"):
                        matrix = _table_rows_to_matrix(ib)
                        if not matrix:
                            continue
                        nrows = len(matrix)
                        ncols = max(len(r) for r in matrix) if matrix else 0
                        if nrows == 0 or ncols == 0:
                            continue
                        tbl = doc.add_table(rows=nrows, cols=ncols)
                        tbl.style = "Table Grid"
                        for ri, row in enumerate(matrix):
                            for ci in range(ncols):
                                val = row[ci] if ci < len(row) else ""
                                tbl.rows[ri].cells[ci].text = val
            else:
                walk(block.get("blocks") or [], list_numbered=list_numbered)

    for page in _sorted_pages(pages):
        walk(page.get("blocks") or [], list_numbered=None)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def emit_pdf_from_pages(
    pages: list[dict[str, Any]],
    image_bytes: dict[str, tuple[bytes, str]],
    out_path: Path,
) -> None:
    try:
        import fitz  # type: ignore
    except ImportError as e:
        raise RuntimeError("pymupdf (fitz) required for pdf export") from e

    doc = fitz.open()
    margin = 56.0
    page_w, page_h = fitz.paper_size("a4")
    page = doc.new_page(width=page_w, height=page_h)
    y = margin
    x0 = margin
    bottom = page_h - margin

    def new_page() -> None:
        nonlocal page, y
        page = doc.new_page(width=page_w, height=page_h)
        y = margin

    def add_lines(text: str, font_size: float) -> None:
        nonlocal page, y
        if not (text or "").strip():
            return
        font = "china-s"
        for line in text.split("\n"):
            if y + font_size * 1.5 > bottom:
                new_page()
            try:
                page.insert_text((x0, y), line, fontsize=font_size, fontname=font)
            except Exception:
                page.insert_text((x0, y), line, fontsize=font_size, fontname="helv")
            y += font_size * 1.35
        y += 4

    def walk(blocks: list, *, heading_scale: float = 1.0) -> None:
        nonlocal page, y
        for block in _sorted_blocks(blocks):
            btype = str(block.get("type") or "")
            if btype == "title":
                t = _plain_text_from_lines(block.get("lines")).strip()
                if t:
                    lvl = max(1, min(3, int(block.get("level") or 1)))
                    fs = (16 - (lvl - 1) * 1.5) * heading_scale
                    if y + fs * 3 > bottom:
                        new_page()
                    add_lines(t, max(10, fs))
            elif btype == "paragraph":
                t = _plain_text_from_lines(block.get("lines")).strip()
                if t:
                    add_lines(t, 11)
            elif btype in ("formula", "code"):
                t = _plain_text_from_lines(block.get("lines")).strip() or str(
                    block.get("formulaLatex") or ""
                ).strip()
                if t:
                    add_lines(t, 10)
            elif btype == "list":
                walk(block.get("blocks") or [], heading_scale=heading_scale)
            elif btype == "list_item":
                t = _plain_text_from_lines(block.get("lines")).strip()
                if t:
                    add_lines(f"• {t}", 11)
                walk(block.get("blocks") or [], heading_scale=heading_scale)
            elif btype == "figure":
                ref = str(block.get("assetRef") or "").strip()
                if ref and ref in image_bytes:
                    raw, _ct = image_bytes[ref]
                    if y + 200 > bottom:
                        new_page()
                    ir = fitz.Rect(x0, y, page_w - margin, y + 220)
                    try:
                        page.insert_image(ir, stream=raw)
                        y = ir.y1 + 12
                    except Exception:
                        add_lines(f"[image failed: {ref}]", 10)
                elif ref:
                    add_lines(f"[missing image: {ref}]", 10)
                walk(block.get("blocks") or [], heading_scale=heading_scale)
            elif btype == "table":
                inner_blocks = block.get("blocks") or []
                for ib in _sorted_blocks(inner_blocks):
                    if not isinstance(ib, dict):
                        continue
                    it = str(ib.get("type") or "")
                    if it == "paragraph":
                        cap = _plain_text_from_lines(ib.get("lines")).strip()
                        if cap:
                            add_lines(cap, 10)
                    elif it == "table" and ib.get("rows"):
                        matrix = _table_rows_to_matrix(ib)
                        for row in matrix:
                            line = " | ".join(cell for cell in row if cell)
                            if line:
                                add_lines(line, 9)
            else:
                walk(block.get("blocks") or [], heading_scale=heading_scale)

    for pg in _sorted_pages(pages):
        walk(pg.get("blocks") or [])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    doc.close()


def emit_xlsx_from_pages(pages: list[dict[str, Any]], out_path: Path) -> None:
    from openpyxl import Workbook  # type: ignore

    wb = Workbook()
    first_sheet = True
    sheet_n = 0

    def harvest_tables(blocks: list) -> None:
        nonlocal first_sheet, sheet_n
        for block in _sorted_blocks(blocks):
            btype = str(block.get("type") or "")
            if btype == "table":
                inner = _find_inner_table_with_rows(block)
                if not inner:
                    harvest_tables(block.get("blocks") or [])
                    continue
                matrix = _table_rows_to_matrix(inner)
                if not matrix:
                    harvest_tables(block.get("blocks") or [])
                    continue
                sheet_n += 1
                title = f"T{sheet_n}"[:31]
                if first_sheet:
                    ws = wb.active
                    ws.title = title
                    first_sheet = False
                else:
                    ws = wb.create_sheet(title=title)
                for ri, row in enumerate(matrix, start=1):
                    for ci, val in enumerate(row, start=1):
                        ws.cell(row=ri, column=ci, value=val)
            harvest_tables(block.get("blocks") or [])

    for pg in _sorted_pages(pages):
        harvest_tables(pg.get("blocks") or [])

    if first_sheet:
        wb.active.append(["(no table blocks in parseDocument)"])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))


def primary_structured_path(out_dir: Path, file_ext: str) -> Path | None:
    ext = (file_ext or "").lower().lstrip(".")
    if ext in ("docx", "doc"):
        return out_dir / "rebuilt.docx"
    if ext == "pdf":
        return out_dir / "rebuilt.pdf"
    if ext in ("xlsx", "xls"):
        return out_dir / "rebuilt.xlsx"
    return None


def emit_structured_by_file_ext(
    pd: dict[str, Any],
    embedded: dict[str, dict[str, Any]],
    out_dir: Path,
) -> dict[str, Any]:
    """Write rebuilt.(docx|pdf|xlsx) from ``pages`` + embedded images. Returns manifest fragment."""
    pages = pd.get("pages") if isinstance(pd.get("pages"), list) else []
    path = primary_structured_path(out_dir, str(pd.get("fileExt") or ""))
    if not path:
        return {"status": "skipped", "reason": "fileExt not doc/docx/pdf/xls/xlsx"}
    image_bytes = build_image_bytes_from_embedded(embedded)
    try:
        if path.suffix.lower() == ".docx":
            emit_docx_from_pages(pages, image_bytes, path)
        elif path.suffix.lower() == ".pdf":
            emit_pdf_from_pages(pages, image_bytes, path)
        elif path.suffix.lower() == ".xlsx":
            emit_xlsx_from_pages(pages, path)
    except Exception as e:
        return {"status": "error", "path": str(path), "error": f"{type(e).__name__}: {e}"}
    return {"status": "ok", "path": str(path)}


def _ensure_chunk_view(pd: dict[str, Any]) -> dict[str, Any]:
    cv = pd.get("chunkView")
    if isinstance(cv, dict) and isinstance(cv.get("blocks"), list):
        return cv
    pages = pd.get("pages")
    if isinstance(pages, list):
        return build_chunk_view(pages)
    return {"blocks": []}


def chunk_view_to_markdown(blocks: list[dict[str, Any]], image_paths: dict[str, str]) -> str:
    lines: list[str] = []
    for blk in blocks:
        if not isinstance(blk, dict):
            continue
        btype = str(blk.get("type") or "")
        if btype == "heading":
            level = max(1, min(6, int(blk.get("level") or 1)))
            text = str(blk.get("text") or "").strip()
            if text:
                lines.append(f"{'#' * level} {text}")
                lines.append("")
        elif btype == "paragraph":
            text = str(blk.get("text") or "").strip()
            if text:
                lines.append(text)
                lines.append("")
        elif btype == "image":
            key = str(blk.get("imageKey") or "").strip()
            rel = image_paths.get(key)
            if rel:
                lines.append(f"![{key}]({rel})")
            else:
                lines.append(f"<!-- image missing payload: {key} -->")
            lines.append("")
        elif btype in ("table_summary", "table_row"):
            text = str(blk.get("text") or blk.get("rowKvText") or "").strip()
            if text:
                lines.append(text)
                lines.append("")
        else:
            text = str(blk.get("text") or "").strip()
            if text:
                lines.append(text)
                lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def chunk_view_to_html(
    title: str,
    blocks: list[dict[str, Any]],
    image_paths: dict[str, str],
) -> str:
    parts: list[str] = [
        "<!DOCTYPE html>",
        "<html lang=\"zh-CN\"><head><meta charset=\"utf-8\">",
        f"<title>{html.escape(title)}</title>",
        "<style>body{font-family:system-ui,sans-serif;max-width:48rem;margin:1rem auto;}"
        "img{max-width:100%;height:auto}</style>",
        "</head><body>",
        f"<h1>{html.escape(title)}</h1>",
    ]
    for blk in blocks:
        if not isinstance(blk, dict):
            continue
        btype = str(blk.get("type") or "")
        if btype == "heading":
            level = max(1, min(6, int(blk.get("level") or 1)))
            text = str(blk.get("text") or "").strip()
            if text:
                parts.append(f"<h{level}>{html.escape(text)}</h{level}>")
        elif btype == "paragraph":
            text = str(blk.get("text") or "").strip()
            if text:
                parts.append(f"<p>{html.escape(text)}</p>")
        elif btype == "image":
            key = str(blk.get("imageKey") or "").strip()
            rel = image_paths.get(key)
            if rel:
                parts.append(f'<p><img src="{html.escape(rel)}" alt="{html.escape(key)}" /></p>')
            else:
                parts.append(f"<!-- missing image: {html.escape(key)} -->")
        elif btype in ("table_summary", "table_row"):
            text = str(blk.get("text") or blk.get("rowKvText") or "").strip()
            if text:
                parts.append(f"<p>{html.escape(text)}</p>")
        else:
            text = str(blk.get("text") or "").strip()
            if text:
                parts.append(f"<p>{html.escape(text)}</p>")
    parts.append("</body></html>")
    return "\n".join(parts) + "\n"


def run(
    input_path: Path,
    out_dir: Path,
    formats: set[str],
    embedded_json: Path | None,
) -> None:
    root = _load_json(input_path)
    if not isinstance(root, dict):
        raise SystemExit("JSON root must be an object")

    pd = _parse_document(root)
    embedded = _embedded_map(root)
    if embedded_json and embedded_json.is_file():
        extra = _load_json(embedded_json)
        if isinstance(extra, dict):
            embedded.update(_embedded_map({"embeddedImagePayloads": extra}))
        else:
            raise SystemExit("--embedded-json must contain an object")

    out_dir.mkdir(parents=True, exist_ok=True)
    assets_dir = out_dir / "assets"
    image_paths: dict[str, str] = {}
    # Always decode ``embeddedImagePayloads`` to files when present, even if user only asked for ``plain``,
    # so the "embedded" half of the response is reconstructed alongside ``parseDocument``.
    if embedded:
        image_paths = extract_images(embedded, assets_dir)

    title = str((pd.get("meta") or {}).get("sourceFileName") or "document").strip() or "document"
    pages = pd.get("pages") if isinstance(pd.get("pages"), list) else []
    cv = _ensure_chunk_view(pd)
    blocks = list(cv.get("blocks") or [])

    if formats & {"plain", "all"}:
        plain = raw_text_from_pages(pages)
        (out_dir / "document.txt").write_text(plain, encoding="utf-8")
    if formats & {"markdown", "all"}:
        md = chunk_view_to_markdown(blocks, image_paths)
        (out_dir / "document.md").write_text(md, encoding="utf-8")
    if formats & {"html", "all"}:
        ht = chunk_view_to_html(title, blocks, image_paths)
        (out_dir / "document.html").write_text(ht, encoding="utf-8")

    structured_info: dict[str, Any] = {}
    if formats & {"structured", "all"}:
        structured_info = emit_structured_by_file_ext(pd, embedded, out_dir)

    manifest = {
        "sourceJson": str(input_path),
        "schemaVersion": pd.get("schemaVersion"),
        "parserKind": pd.get("parserKind"),
        "fileExt": pd.get("fileExt"),
        "imagesWritten": sorted(image_paths.keys()),
        "structured": structured_info,
        "note": "Semantic re-export from parseDocument + embeddedImagePayloads; not original binary.",
    }
    (out_dir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("-i", "--input", type=Path, required=True, help="Parse JSON file")
    ap.add_argument("-o", "--out", type=Path, default=Path("reconstructed"), help="Output directory")
    ap.add_argument(
        "--embedded-json",
        type=Path,
        default=None,
        help="Optional JSON object of {imageKey: {base64, contentType, ...}}",
    )
    ap.add_argument(
        "--formats",
        default="all",
        help="Comma-separated: plain, markdown, html, images, structured, all (default: all)",
    )
    args = ap.parse_args()
    parts = {p.strip().lower() for p in args.formats.split(",") if p.strip()}
    allowed = {"plain", "markdown", "html", "images", "structured", "all"}
    if not parts <= allowed:
        raise SystemExit(f"--formats must be subset of {sorted(allowed)}")
    if not parts:
        parts = {"all"}
    run(args.input, args.out, parts, args.embedded_json)


if __name__ == "__main__":
    main()
