"""DOCX native extended parse: header/footer, meta.readingOrderPolicy, sourceRegion."""

from __future__ import annotations

import io

import pytest

pytest.importorskip("docx")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from app.parsers.docx_parser import DocxParser, _blocks_from_notes_xml
from app.parsers.unified_builders import IdGen, raw_text_from_pages


def _add_external_hyperlink(paragraph: object, url: str, text: str) -> None:
    part = paragraph.part  # type: ignore[attr-defined]
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    h.append(r)
    paragraph._p.append(h)  # type: ignore[attr-defined]


def test_docx_header_footer_body_and_chunkview_source_region() -> None:
    buf = io.BytesIO()
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "HDR-UNIQUE-42"
    doc.sections[0].footer.paragraphs[0].text = "FTR-UNIQUE-43"
    doc.add_paragraph("Main body text")
    doc.save(buf)
    parser = DocxParser()
    result = parser.parse(buf.getvalue(), filename="t.docx", file_ext="docx", mime_type="")
    assert "HDR-UNIQUE-42" in result.raw_text
    assert "FTR-UNIQUE-43" in result.raw_text
    assert "Main body text" in result.raw_text
    pd = result.parse_document or {}
    assert pd.get("parserKind") == "docx_python_docx_v3"
    assert pd.get("meta", {}).get("readingOrderPolicy")
    blocks = pd["pages"][0]["blocks"]
    regions = [b.get("sourceRegion") for b in blocks if isinstance(b, dict)]
    assert "header" in regions
    assert "footer" in regions
    cv = pd.get("chunkView") or {}
    cv_blocks = cv.get("blocks") or []
    hdr_chunks = [b for b in cv_blocks if b.get("sourceRegion") == "header"]
    assert hdr_chunks
    assert "HDR-UNIQUE-42" in hdr_chunks[0].get("text", "")


def test_docx_hyperlink_visible_text_in_raw_and_blocks() -> None:
    buf = io.BytesIO()
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("LEFT-")
    _add_external_hyperlink(p, "https://example.com/", "MID-HLINK")
    p.add_run("-RIGHT")
    doc.save(buf)
    parser = DocxParser()
    result = parser.parse(buf.getvalue(), filename="t.docx", file_ext="docx", mime_type="")
    assert "LEFT-" in result.raw_text
    assert "MID-HLINK" in result.raw_text
    assert "-RIGHT" in result.raw_text
    pd = result.parse_document or {}
    # Blocks use lines/spans (no top-level block["text"]); flatten like the parser does for raw_text.
    flat_from_pages = raw_text_from_pages(pd.get("pages") or [])
    assert "MID-HLINK" in flat_from_pages


def test_docx_table_cell_hyperlink_and_field_style_extraction() -> None:
    buf = io.BytesIO()
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    p = c.paragraphs[0]
    p.add_run("CELL-A-")
    _add_external_hyperlink(p, "https://example.com/table", "CELL-LINK")
    p.add_run("-CELL-B")
    doc.save(buf)
    parser = DocxParser()
    result = parser.parse(buf.getvalue(), filename="t.docx", file_ext="docx", mime_type="")
    assert "CELL-LINK" in result.raw_text
    pd = result.parse_document or {}
    blocks = (pd.get("pages") or [{}])[0].get("blocks") or []
    tbl = next((b for b in blocks if isinstance(b, dict) and b.get("type") == "table"), None)
    assert tbl is not None
    inner = (tbl.get("blocks") or [{}])[0]
    row0 = (inner.get("rows") or [{}])[0]
    cell0 = (row0.get("cells") or [{}])[0]
    spans = (cell0.get("spans") or [])
    cell_text = "".join((s.get("text") or "") for s in spans if isinstance(s, dict))
    assert "CELL-LINK" in cell_text
    assert "CELL-A-" in cell_text
    assert "-CELL-B" in cell_text


def test_docx_sdt_block_level_body_content() -> None:
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("before-sdt")
    sdt = parse_xml(
        "<w:sdt %s><w:sdtContent><w:p><w:r><w:t>SDT-BLOCK-99</w:t></w:r></w:p></w:sdtContent></w:sdt>"
        % nsdecls("w")
    )
    doc.element.body.append(sdt)
    doc.add_paragraph("after-sdt")
    doc.save(buf)
    parser = DocxParser()
    result = parser.parse(buf.getvalue(), filename="t.docx", file_ext="docx", mime_type="")
    assert "SDT-BLOCK-99" in result.raw_text
    assert "before-sdt" in result.raw_text
    assert "after-sdt" in result.raw_text


def test_footnote_xml_extracts_text_block() -> None:
    xml = (
        b'<?xml version="1.0" encoding="UTF-8"?>'
        b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:footnote w:id="0"><w:p><w:r><w:t>Note-A</w:t></w:r></w:p></w:footnote>'
        b"</w:footnotes>"
    )
    id_gen = IdGen("n")
    seq = iter(range(500))

    def nr() -> int:
        return next(seq)

    blocks = _blocks_from_notes_xml(xml, id_gen=id_gen, next_ro=nr, source_region="footnote")
    assert len(blocks) == 1
    assert blocks[0].get("sourceRegion") == "footnote"
    lines = blocks[0].get("lines") or []
    assert lines and "Note-A" in (lines[0].get("spans") or [{}])[0].get("text", "")
