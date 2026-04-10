"""Structured cleanup: header/footer removal, assets, headingPath."""

from __future__ import annotations

import io

import pytest

from app.services.clean_structured import (
    annotate_chunk_view_heading_paths,
    clean_parse_document,
)


def test_annotate_heading_path_nested_sections() -> None:
    blocks = annotate_chunk_view_heading_paths(
        [
            {"type": "paragraph", "text": "intro"},
            {"type": "heading", "level": 1, "text": "A"},
            {"type": "paragraph", "text": "under A"},
            {"type": "heading", "level": 2, "text": "A.1"},
            {"type": "paragraph", "text": "under A.1"},
        ]
    )
    assert blocks[0]["headingPath"] == []
    assert blocks[1]["headingPath"] == ["A"]
    assert blocks[2]["headingPath"] == ["A"]
    assert blocks[3]["headingPath"] == ["A", "A.1"]
    assert blocks[4]["headingPath"] == ["A", "A.1"]


def test_clean_docx_drops_header_footer_text_and_keeps_body() -> None:
    pytest.importorskip("docx")
    from docx import Document

    from app.parsers.docx_parser import DocxParser

    buf = io.BytesIO()
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "HDR-X"
    doc.sections[0].footer.paragraphs[0].text = "FTR-Y"
    doc.add_paragraph("BODY-Z")
    doc.save(buf)
    parser = DocxParser()
    result = parser.parse(buf.getvalue(), filename="t.docx", file_ext="docx", mime_type="")
    pd = result.parse_document or {}
    cleaned = clean_parse_document(pd)
    texts = " ".join(str(b.get("text") or "") for b in cleaned["chunkView"]["blocks"])
    assert "HDR-X" not in texts and "FTR-Y" not in texts
    assert "BODY-Z" in texts
    assert "header" in cleaned["meta"]["cleaning"]["droppedSourceRegions"]
    assert all("headingPath" in b for b in cleaned["chunkView"]["blocks"])


def test_clean_parse_minimal_dict_filters_regions() -> None:
    pd = {
        "schemaVersion": "2.0",
        "parseRoute": "native",
        "fileExt": "docx",
        "parserKind": "test",
        "meta": {"pageCount": 1, "sourceFileName": "x", "warnings": [], "hasBbox": False},
        "assets": [],
        "pages": [
            {
                "pageIndex": 0,
                "blocks": [
                    {
                        "id": "b1",
                        "type": "paragraph",
                        "readingOrder": 0,
                        "lines": [{"id": "l1", "spans": [{"id": "s1", "text": "drop-me"}]}],
                        "sourceRegion": "header",
                    },
                    {
                        "id": "b2",
                        "type": "paragraph",
                        "readingOrder": 1,
                        "lines": [{"id": "l2", "spans": [{"id": "s2", "text": "keep-me"}]}],
                    },
                ],
            }
        ],
    }
    cleaned = clean_parse_document(pd)
    texts = [str(b.get("text") or "") for b in cleaned["chunkView"]["blocks"]]
    assert "drop-me" not in texts
    assert any("keep-me" in t for t in texts)
